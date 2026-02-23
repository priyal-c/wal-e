"""DOCX Remediation Guide Reporter.

Generates a detailed Word document with step-by-step remediation instructions
for every issue and anti-pattern reported by WAL-E, including cloud-specific
references to public Databricks documentation.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from wal_e.reporters.base import (
    PILLAR_DISPLAY_NAMES,
    PILLAR_ORDER,
    AuditEntry,
    BaseReporter,
    ScoredAssessment,
)

# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
_DATABRICKS_RED = RGBColor(0xFF, 0x38, 0x21)
_GREEN = RGBColor(0x2E, 0x7D, 0x32)
_AMBER = RGBColor(0xF5, 0x7C, 0x00)
_RED = RGBColor(0xC6, 0x28, 0x28)
_BLUE = RGBColor(0x15, 0x65, 0xC0)
_GRAY = RGBColor(0x61, 0x61, 0x61)
_DARK = RGBColor(0x21, 0x21, 0x21)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT_GREEN_BG = "E8F5E9"
_LIGHT_RED_BG = "FFEBEE"
_LIGHT_AMBER_BG = "FFF3E0"
_LIGHT_BLUE_BG = "E3F2FD"
_HEADER_BG = "1565C0"

# ---------------------------------------------------------------------------
# Databricks documentation base URLs
# ---------------------------------------------------------------------------
_DOCS_BASE = {
    "aws": "https://docs.databricks.com/aws/en",
    "azure": "https://docs.databricks.com/azure/en",
    "gcp": "https://docs.databricks.com/gcp/en",
}

def _docs_url(cloud: str, path: str) -> str:
    base = _DOCS_BASE.get(cloud, _DOCS_BASE["aws"])
    return f"{base}/{path.lstrip('/')}"


# ---------------------------------------------------------------------------
# Remediation Knowledge Base
# ---------------------------------------------------------------------------
# Maps pillar/topic → remediation entry with:
#   - why: why this matters
#   - steps: list of remediation steps
#   - docs: dict of {cloud: [(label, doc_path), ...]}
#   - applies_to: list of BP name substrings that this entry covers
# ---------------------------------------------------------------------------

_REMEDIATION_KB: list[dict[str, Any]] = [
    # ── Governance ──
    {
        "topic": "Unity Catalog Adoption",
        "applies_to": ["governance process", "metadata in one place", "centralize access"],
        "why": "Unity Catalog provides a unified governance layer for all data and AI assets. Without it, access control is fragmented across workspace-level ACLs.",
        "steps": [
            "Enable Unity Catalog at the account level and assign a metastore to each workspace.",
            "Migrate existing hive_metastore tables to Unity Catalog using the UCX migration tool.",
            "Create a catalog naming convention (e.g., prod_*, dev_*, staging_*).",
            "Define ownership at the catalog and schema level.",
            "Enable automatic lineage tracking by ensuring all workloads run through UC-enabled compute.",
        ],
        "docs": {
            "all": [
                ("Unity Catalog Setup", "/data-governance/unity-catalog/get-started"),
                ("Migration Guide (UCX)", "/data-governance/unity-catalog/migrate"),
                ("Best Practices", "/data-governance/unity-catalog/best-practices"),
            ],
        },
    },
    {
        "topic": "Data Quality with DLT Expectations",
        "applies_to": ["dq standards", "dq tools", "data quality"],
        "why": "Without data quality checks, bad data propagates downstream silently, causing incorrect analytics and ML model drift.",
        "steps": [
            "Add DLT expectations (EXPECT, EXPECT OR DROP, EXPECT OR FAIL) to all DLT pipelines.",
            "Define quality rules for critical columns: NOT NULL, valid ranges, referential integrity.",
            "Set up DLT event log monitoring to track quality metrics over time.",
            "Create alerts on quality metric regressions.",
        ],
        "docs": {
            "all": [
                ("DLT Expectations", "/delta-live-tables/expectations"),
                ("Monitor Data Quality", "/delta-live-tables/observability"),
            ],
        },
    },
    {
        "topic": "Audit Logging",
        "applies_to": ["audit logging", "audit events", "configure audit"],
        "why": "Audit logs are essential for security incident investigation, compliance reporting, and understanding who accessed what data.",
        "steps": [
            "Enable audit log delivery at the account level.",
            "Configure delivery to a cloud storage bucket (S3/ADLS/GCS).",
            "Set up system tables (system.access.audit) for SQL-based audit analysis.",
            "Create dashboards for login patterns, permission changes, and data access.",
        ],
        "docs": {
            "aws": [("Audit Logging (AWS)", "/administration-guide/account-settings/audit-logs"), ("System Tables", "/administration-guide/system-tables")],
            "azure": [("Audit Logging (Azure)", "/administration-guide/account-settings/azure-diagnostic-logs"), ("System Tables", "/administration-guide/system-tables")],
            "gcp": [("Audit Logging (GCP)", "/administration-guide/account-settings/audit-logs"), ("System Tables", "/administration-guide/system-tables")],
        },
    },
    {
        "topic": "Group Management and BROWSE Privilege",
        "applies_to": ["group management", "browse privilege", "account-level group"],
        "why": "IdP-synced groups ensure permissions stay in sync with your identity provider. BROWSE privilege enables data discovery without granting data access.",
        "steps": [
            "Configure SCIM provisioning from your IdP (Entra ID, Okta, OneLogin) to sync groups automatically.",
            "Create account-level groups instead of workspace-level groups.",
            "Grant BROWSE privilege on catalogs to the 'account users' group for discovery.",
            "Remove workspace-local groups and migrate to account-level groups.",
        ],
        "docs": {
            "all": [
                ("SCIM Provisioning", "/administration-guide/users-groups/scim"),
                ("Account Groups", "/administration-guide/users-groups/groups"),
                ("BROWSE Privilege", "/data-governance/unity-catalog/manage-privileges/privileges#browse"),
            ],
        },
    },
    {
        "topic": "Managed Tables and External Locations",
        "applies_to": ["managed tables", "external location"],
        "why": "Managed tables in Unity Catalog provide automatic lifecycle management, governance, and simplified access control. External tables require manual credential and path management.",
        "steps": [
            "For new tables, always create as managed tables in Unity Catalog.",
            "Migrate existing external tables using CREATE TABLE AS SELECT or DEEP CLONE.",
            "Limit external locations to data that must remain in existing storage (e.g., shared with non-Databricks systems).",
            "Audit and minimize the number of external locations and storage credentials.",
        ],
        "docs": {
            "all": [
                ("Managed vs External Tables", "/data-governance/unity-catalog/create-tables"),
                ("External Locations", "/data-governance/unity-catalog/manage-external-locations"),
            ],
        },
    },
    # ── Security ──
    {
        "topic": "SSO and SCIM Configuration",
        "applies_to": ["sso config", "scim provisioning", "sso"],
        "why": "Without SSO, users manage separate passwords for Databricks. Without SCIM, group membership drifts from your identity provider, creating orphaned access.",
        "steps": [
            "Configure SSO at the account level with your identity provider.",
            "Enable SCIM provisioning to auto-sync users and groups.",
            "Disable local password authentication after SSO is verified.",
            "Set up SCIM for all workspaces in the account.",
        ],
        "docs": {
            "aws": [("SSO Setup (AWS)", "/administration-guide/users-groups/single-sign-on"), ("SCIM (AWS)", "/administration-guide/users-groups/scim/aad")],
            "azure": [("SSO Setup (Azure)", "/administration-guide/users-groups/single-sign-on"), ("SCIM (Azure)", "/administration-guide/users-groups/scim/aad")],
            "gcp": [("SSO Setup (GCP)", "/administration-guide/users-groups/single-sign-on"), ("SCIM (GCP)", "/administration-guide/users-groups/scim")],
        },
    },
    {
        "topic": "Network Security and VPC/VNET",
        "applies_to": ["network security", "customer-managed vpc", "ip access", "vpc"],
        "why": "Without network isolation, the Databricks control plane communicates over public internet. A customer-managed VPC/VNET restricts data plane traffic to your network.",
        "steps": [
            "Deploy Databricks in a customer-managed VPC (AWS), VNET (Azure), or use Private Service Connect (GCP).",
            "Enable IP access lists to restrict workspace access to known corporate IPs/VPNs.",
            "Configure Private Link / Private Endpoints for control plane and storage connectivity.",
            "Disable public access to the workspace if all users connect via VPN.",
        ],
        "docs": {
            "aws": [("Customer-Managed VPC", "/administration-guide/cloud-configurations/aws/customer-managed-vpc"), ("Private Link", "/administration-guide/cloud-configurations/aws/privatelink"), ("IP Access Lists", "/security/network/front-end/ip-access-list")],
            "azure": [("VNET Injection", "/administration-guide/cloud-configurations/azure/vnet-inject"), ("Private Link", "/administration-guide/cloud-configurations/azure/private-link"), ("IP Access Lists", "/security/network/front-end/ip-access-list")],
            "gcp": [("PSC", "/administration-guide/cloud-configurations/gcp/private-service-connect"), ("IP Access Lists", "/security/network/front-end/ip-access-list")],
        },
    },
    {
        "topic": "Service Principals for Automation",
        "applies_to": ["service principal", "service principals for automation"],
        "why": "Using personal accounts for automated jobs creates single points of failure and security risk. Service principals provide auditable, role-based automation identities.",
        "steps": [
            "Create service principals for each automation domain (ETL, ML, CI/CD).",
            "Transfer job ownership from personal accounts to service principals.",
            "Use OAuth tokens or managed identity instead of PAT tokens for service principals.",
            "Apply least-privilege permissions to each service principal.",
        ],
        "docs": {
            "all": [
                ("Service Principals", "/administration-guide/users-groups/service-principals"),
                ("OAuth M2M", "/dev-tools/auth/oauth-m2m"),
            ],
            "azure": [("Managed Identity", "/administration-guide/users-groups/service-principals#managed-identity")],
        },
    },
    {
        "topic": "DBFS Root Restriction",
        "applies_to": ["dbfs root", "restrict dbfs"],
        "why": "DBFS root is a shared, unmanaged storage location. Sensitive data stored there lacks fine-grained access controls and is accessible to all workspace users.",
        "steps": [
            "Disable the DBFS browser in workspace settings (enableDbfsFileBrowser = false).",
            "Migrate data from dbfs:/user/ and dbfs:/FileStore/ to Unity Catalog managed volumes.",
            "Use workspace settings to block notebook export if data exfiltration is a concern.",
            "Audit existing DBFS usage and create a migration plan.",
        ],
        "docs": {
            "all": [
                ("DBFS Best Practices", "/dbfs/best-practices"),
                ("Unity Catalog Volumes", "/data-governance/unity-catalog/create-volumes"),
            ],
        },
    },
    {
        "topic": "Failed Login and Permission Change Monitoring",
        "applies_to": ["failed login", "permission change audit"],
        "why": "High numbers of failed logins may indicate brute-force attacks. Frequent permission changes without governance suggest access control drift.",
        "steps": [
            "Enable system tables (system.access.audit) and query for failed login patterns.",
            "Set up SQL alerts for >10 failed logins in 1 hour from the same IP.",
            "Create a dashboard tracking permission change events over time.",
            "Establish a change approval process for permission modifications.",
        ],
        "docs": {
            "all": [
                ("Audit Log System Table", "/administration-guide/system-tables/audit-logs"),
                ("SQL Alerts", "/sql/user/alerts"),
            ],
        },
    },
    # ── Compute / Performance ──
    {
        "topic": "Cluster Policies",
        "applies_to": ["predefined compute", "standardize compute", "cluster policies cost", "compute templates"],
        "why": "Without cluster policies, users can create oversized or misconfigured clusters, leading to cost overruns and security gaps.",
        "steps": [
            "Create cluster policies for each team/use case (dev, production, ML training).",
            "Restrict instance types, max workers, and auto-termination in policies.",
            "Set default DBR version and enforce Unity Catalog access mode.",
            "Assign policies to groups, not individual users.",
        ],
        "docs": {
            "all": [
                ("Cluster Policies", "/administration-guide/clusters/policies"),
                ("Policy Best Practices", "/administration-guide/clusters/policy-best-practices"),
            ],
        },
    },
    {
        "topic": "Auto-Termination and Idle Cluster Waste",
        "applies_to": ["auto-termination", "idle cluster", "restart long-running", "always-on"],
        "why": "Clusters left running 24/7 without workloads waste significant DBUs. Auto-termination ensures clusters shut down after a period of inactivity.",
        "steps": [
            "Set auto-termination to 10-20 minutes for interactive clusters.",
            "Use job clusters (ephemeral) instead of all-purpose clusters for scheduled workloads.",
            "Enforce auto-termination through cluster policies (cannot be disabled by users).",
            "Review system.compute.clusters for clusters with >500 running hours/month.",
        ],
        "docs": {
            "all": [
                ("Auto-Termination", "/compute/configure#auto-termination"),
                ("Job Clusters", "/jobs/use-compute"),
            ],
        },
    },
    {
        "topic": "Serverless Compute and SQL Warehouses",
        "applies_to": ["serverless", "sql for sql", "serverless compute"],
        "why": "Serverless compute eliminates infrastructure management, provides instant startup, and automatically scales. SQL warehouses are purpose-built for BI/SQL workloads.",
        "steps": [
            "Enable serverless SQL warehouses for all BI and ad-hoc SQL workloads.",
            "Migrate from classic compute to serverless where supported.",
            "Use serverless jobs for ETL and scheduled workloads.",
            "Set auto-stop on warehouses (5-10 minute inactivity).",
        ],
        "docs": {
            "all": [
                ("Serverless SQL Warehouses", "/compute/sql-warehouse/serverless"),
                ("Serverless Compute", "/compute/serverless"),
            ],
        },
    },
    {
        "topic": "Instance Type Optimization",
        "applies_to": ["graviton", "right instance", "instance type", "standard access mode"],
        "why": "Choosing the right instance type can reduce costs by 20-40%. Graviton (AWS), Dv5/Ev5 (Azure), and T2D (GCP) offer better price-performance than older generations.",
        "steps": [
            "Audit current instance types across all clusters.",
            "Migrate to latest-generation instances: Graviton (AWS), Dv5/Ev5 (Azure), T2D/N2D (GCP).",
            "Use memory-optimized instances for Spark shuffle-heavy workloads.",
            "Use compute-optimized instances for ML training.",
            "Enforce via cluster policies.",
        ],
        "docs": {
            "aws": [("Graviton Support", "/compute/graviton"), ("Instance Types", "/compute/supported-instance-types")],
            "azure": [("VM Types", "/compute/supported-instance-types")],
            "gcp": [("Machine Types", "/compute/supported-instance-types")],
        },
    },
    {
        "topic": "Query Performance and Slow Queries",
        "applies_to": ["query failure", "slow query", "warehouse utilization"],
        "why": "High query failure rates and slow queries indicate performance bottlenecks that impact user productivity and increase costs.",
        "steps": [
            "Query system.query.history to identify the top 20 slowest queries by P95 duration.",
            "Check for missing optimizations: Z-ORDER/liquid clustering, file compaction, caching.",
            "Right-size SQL warehouses based on utilization (scale up for concurrency, not just size).",
            "Use Query Profile to diagnose individual slow queries.",
            "Enable predictive optimization for automated OPTIMIZE and VACUUM.",
        ],
        "docs": {
            "all": [
                ("Query Profile", "/sql/user/queries/query-profile"),
                ("Predictive Optimization", "/delta/predictive-optimization"),
                ("Query History System Table", "/administration-guide/system-tables/query-history"),
            ],
        },
    },
    {
        "topic": "Liquid Clustering and Predictive Optimization",
        "applies_to": ["liquid clustering", "predictive optimization"],
        "why": "Traditional partitioning and Z-ORDER require manual tuning. Liquid clustering adapts automatically, and predictive optimization runs OPTIMIZE/VACUUM without manual scheduling.",
        "steps": [
            "Enable predictive optimization at the catalog or schema level.",
            "Migrate tables from static partitioning to liquid clustering (ALTER TABLE ... CLUSTER BY).",
            "Remove manual OPTIMIZE and VACUUM jobs — predictive optimization handles this.",
            "Monitor via system tables to verify optimization is running.",
        ],
        "docs": {
            "all": [
                ("Liquid Clustering", "/delta/clustering"),
                ("Predictive Optimization", "/delta/predictive-optimization"),
            ],
        },
    },
    # ── Operations ──
    {
        "topic": "CI/CD and Git Integration",
        "applies_to": ["standardize ci/cd", "enterprise scm", "git", "iac"],
        "why": "Without CI/CD, code changes are deployed manually, increasing risk of errors and making rollbacks difficult.",
        "steps": [
            "Connect Databricks Repos to your Git provider (GitHub, Azure DevOps, GitLab).",
            "Use Databricks Asset Bundles (DABs) for deploying jobs, pipelines, and ML models.",
            "Implement a multi-branch workflow: dev → staging → production.",
            "Use Terraform or Pulumi for workspace infrastructure provisioning.",
            "Set up automated testing in CI pipelines before deployment.",
        ],
        "docs": {
            "all": [
                ("Databricks Repos", "/repos"),
                ("Asset Bundles (DABs)", "/dev-tools/bundles"),
                ("CI/CD Patterns", "/dev-tools/ci-cd"),
                ("Terraform Provider", "/dev-tools/terraform"),
            ],
        },
    },
    {
        "topic": "Job Retries and Resilience",
        "applies_to": ["auto retries", "job automation recovery", "job success rate", "recurring job failure"],
        "why": "Transient failures (network hiccups, spot instance reclamation) are normal in cloud. Without retries, jobs fail unnecessarily and require manual intervention.",
        "steps": [
            "Set max_retries to 1-3 on all production jobs.",
            "Configure retry delays with exponential backoff for API-dependent jobs.",
            "Use task-level retries (not just job-level) for multi-task workflows.",
            "Set up email/Slack alerts on job failures.",
            "Review system.lakeflow.job_run_timeline for chronic failures.",
        ],
        "docs": {
            "all": [
                ("Job Retries", "/jobs/settings#retries"),
                ("Task Orchestration", "/jobs/use-task-orchestration"),
                ("Job Run System Table", "/administration-guide/system-tables/lakeflow"),
            ],
        },
    },
    {
        "topic": "Cluster Utilization Efficiency",
        "applies_to": ["cluster utilization efficiency", "24/7"],
        "why": "Clusters running 24/7 without workloads are the single largest source of wasted compute spend.",
        "steps": [
            "Query system.compute.clusters to find clusters with >500 running hours/month.",
            "Convert always-on clusters to job clusters (ephemeral, created per job run).",
            "For interactive use cases, set aggressive auto-termination (10 min).",
            "Consider serverless compute for unpredictable workloads.",
        ],
        "docs": {
            "all": [
                ("Job Clusters", "/jobs/use-compute"),
                ("Compute System Table", "/administration-guide/system-tables/compute"),
            ],
        },
    },
    # ── Reliability ──
    {
        "topic": "Auto-Scaling for ETL and DLT",
        "applies_to": ["etl autoscaling", "sql warehouse autoscaling", "auto-scaling", "scaling"],
        "why": "Fixed-size clusters either waste resources during low load or become bottlenecks during peak. Auto-scaling adjusts dynamically.",
        "steps": [
            "Enable autoscale on all-purpose and job clusters (min_workers < max_workers).",
            "Enable enhanced autoscaling on DLT pipelines.",
            "Set SQL warehouse min_num_clusters=1, max_num_clusters based on concurrency needs.",
            "Monitor autoscaling behavior via cluster event logs.",
        ],
        "docs": {
            "all": [
                ("Cluster Autoscaling", "/compute/configure#autoscaling"),
                ("DLT Enhanced Autoscaling", "/delta-live-tables/settings#enhanced-autoscaling"),
                ("SQL Warehouse Scaling", "/compute/sql-warehouse/create#scaling"),
            ],
        },
    },
    # ── Cost ──
    {
        "topic": "Spot/Preemptible Instance Strategy",
        "applies_to": ["spot instance", "spot", "preemptible"],
        "why": "Spot instances (AWS), Spot VMs (Azure), and Preemptible VMs (GCP) offer 60-90% cost savings for fault-tolerant batch workloads.",
        "steps": [
            "Use spot instances for worker nodes on batch/ETL job clusters.",
            "Keep the driver node on-demand (first_on_demand=1) for reliability.",
            "Set spot_bid_price_percent to 100% for best availability.",
            "Avoid spot for interactive/streaming clusters where interruptions are costly.",
            "Configure fallback to on-demand when spot capacity is unavailable.",
        ],
        "docs": {
            "aws": [("Spot Instances", "/compute/configure#spot-instances")],
            "azure": [("Azure Spot VMs", "/compute/configure#spot-instances")],
            "gcp": [("Preemptible VMs", "/compute/configure#spot-instances")],
        },
    },
    {
        "topic": "Cost Monitoring and Budget Alerts",
        "applies_to": ["monitor costs", "cost reports", "budget alert", "cost trend", "concentration risk"],
        "why": "Without cost monitoring, spend grows silently. Budget alerts provide early warning before overruns become critical.",
        "steps": [
            "Enable system.billing.usage system table.",
            "Create a billing dashboard showing daily/weekly DBU trends by SKU and cluster.",
            "Set up budget alerts in the Databricks Account Console.",
            "Implement chargeback by tagging clusters with team/project/cost-center.",
            "Review top-cost clusters monthly and right-size or terminate.",
        ],
        "docs": {
            "all": [
                ("Billing System Table", "/administration-guide/system-tables/billing"),
                ("Budget Alerts", "/administration-guide/account-settings/budgets"),
                ("Cluster Tags", "/compute/configure#cluster-tags"),
            ],
            "aws": [("AWS Cost Explorer", "/administration-guide/account-settings/usage-analysis")],
            "azure": [("Azure Cost Management", "/administration-guide/account-settings/usage-analysis")],
        },
    },
    {
        "topic": "Reserved Instances and Committed Use",
        "applies_to": ["on-demand vs reserved", "reserved"],
        "why": "For steady-state workloads that run predictably, reserved capacity (Savings Plans, Reserved Instances, CUDs) can reduce costs by 30-60%.",
        "steps": [
            "Analyze system.billing.usage for consistent daily DBU usage patterns.",
            "Identify baseline (always-on) DBU consumption vs. burst usage.",
            "Purchase Databricks Committed Use Discounts for the baseline.",
            "Layer cloud-provider reserved instances for the underlying VMs.",
        ],
        "docs": {
            "aws": [("Savings Plans", "/administration-guide/account-settings/usage-analysis")],
            "azure": [("Reserved Instances", "/administration-guide/account-settings/usage-analysis")],
            "gcp": [("Committed Use Discounts", "/administration-guide/account-settings/usage-analysis")],
        },
    },
    {
        "topic": "Job Clusters for Cost Efficiency",
        "applies_to": ["job cluster"],
        "why": "All-purpose clusters persist between job runs, accruing idle costs. Job clusters are created on demand and terminated immediately after the job completes.",
        "steps": [
            "For every scheduled job, switch from 'existing cluster' to 'new job cluster'.",
            "Define job cluster specs in the job definition (instance type, workers, DBR version).",
            "Use cluster policies to standardize job cluster configurations.",
            "Enable spot instances on job cluster worker nodes for batch workloads.",
        ],
        "docs": {
            "all": [
                ("Job Compute", "/jobs/use-compute"),
                ("Job Cluster Configuration", "/jobs/settings#compute"),
            ],
        },
    },
    {
        "topic": "Tagging for Chargeback",
        "applies_to": ["tag cluster", "chargeback"],
        "why": "Without tagging, you cannot attribute costs to teams, projects, or environments. Tags enable chargeback and showback reporting.",
        "steps": [
            "Define mandatory tags: team, project, environment, cost-center.",
            "Enforce tags via cluster policies (custom_tags required).",
            "Query system.billing.usage with tag filters for chargeback reports.",
            "Create team-level billing dashboards.",
        ],
        "docs": {
            "all": [
                ("Cluster Tags", "/compute/configure#cluster-tags"),
                ("Tag-Based Billing", "/administration-guide/account-settings/usage-analysis"),
            ],
        },
    },
]


def _find_remediation(bp_name: str) -> dict[str, Any] | None:
    """Find the best matching remediation entry for a best practice name."""
    name_lower = bp_name.lower()
    for entry in _REMEDIATION_KB:
        for pattern in entry["applies_to"]:
            if pattern.lower() in name_lower:
                return entry
    return None


def _get_docs_for_cloud(entry: dict[str, Any], cloud: str) -> list[tuple[str, str]]:
    """Get documentation links for a specific cloud from a remediation entry."""
    docs = entry.get("docs", {})
    result: list[tuple[str, str]] = []
    for label, path in docs.get("all", []):
        result.append((label, _docs_url(cloud, path)))
    for label, path in docs.get(cloud, []):
        result.append((label, _docs_url(cloud, path)))
    return result


# ---------------------------------------------------------------------------
# Cell shading helper
# ---------------------------------------------------------------------------
def _shade_cell(cell: Any, hex_color: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    tc_pr.append(shading)


# ---------------------------------------------------------------------------
# DOCX Remediation Reporter
# ---------------------------------------------------------------------------

class DocxRemediationReporter(BaseReporter):
    """Generates a Word document with detailed remediation instructions."""

    def __init__(self) -> None:
        super().__init__("WAL_Assessment_Remediation_Guide.docx")

    def generate(
        self,
        scored_assessment: ScoredAssessment,
        collected_data: Dict[str, Any],
        audit_entries: List[AuditEntry],
        output_dir: Union[str, Path],
    ) -> Path:
        output_path = self._ensure_output_dir(output_dir) / self.output_filename
        cloud = self._get_cloud_provider(scored_assessment)
        cloud_display = self._cloud_display_name(cloud)
        cloud_short = self._cloud_short_name(cloud)
        workspace = self._get_workspace_host(scored_assessment)
        date = self._get_assessment_date(scored_assessment)
        maturity = self._get_maturity_level(scored_assessment)
        overall = self._get_overall_score(scored_assessment)
        verified = scored_assessment.get("verified_score", overall)
        coverage = scored_assessment.get("coverage_pct", 100.0)
        bps = self._get_best_practice_scores(scored_assessment)

        doc = Document()

        self._set_default_font(doc)
        self._add_title_page(doc, workspace, cloud_display, cloud_short, date, maturity, verified, coverage)
        self._add_executive_summary(doc, bps, cloud_short, verified, coverage)

        for pillar in PILLAR_ORDER:
            pillar_bps = self._get_bps_for_pillar(bps, pillar)
            if not pillar_bps:
                continue
            issues = [bp for bp in pillar_bps if bp.get("score", 2) < 2]
            if not issues:
                continue
            self._add_pillar_section(doc, pillar, issues, cloud, cloud_short)

        self._add_appendix(doc, bps, cloud_short)

        doc.save(str(output_path))
        return output_path

    def _set_default_font(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)
        font.color.rgb = _DARK
        for level in range(1, 5):
            h = doc.styles[f"Heading {level}"]
            h.font.name = "Calibri"
            h.font.color.rgb = _BLUE

    def _add_title_page(self, doc: Document, workspace: str, cloud_display: str,
                        cloud_short: str, date: str, maturity: str,
                        verified: float, coverage: float) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n\n\n")
        run.font.size = Pt(24)

        title = doc.add_heading("WAL-E Remediation Guide", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Well-Architected Lakehouse Assessment\nDetailed Remediation Instructions")
        run.font.size = Pt(14)
        run.font.color.rgb = _GRAY

        doc.add_paragraph()

        table = doc.add_table(rows=6, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta = [
            ("Workspace", workspace),
            ("Cloud Provider", f"{cloud_display} ({cloud_short})"),
            ("Assessment Date", date),
            ("Maturity Level", maturity),
            ("Verified Score", f"{(verified / 2.0) * 100:.0f}%"),
            ("Assessment Coverage", f"{coverage:.0f}%"),
        ]
        for i, (label, value) in enumerate(meta):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = value
            for cell in cells:
                for par in cell.paragraphs:
                    par.style = doc.styles["Normal"]
            cells[0].paragraphs[0].runs[0].font.bold = True
            cells[0].paragraphs[0].runs[0].font.color.rgb = _BLUE

        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, bps: list, cloud_short: str,
                               verified: float, coverage: float) -> None:
        doc.add_heading("Executive Summary", level=1)

        critical = [bp for bp in bps if bp.get("score", 2) == 0 and bp.get("verified", True)]
        partial = [bp for bp in bps if bp.get("score", 2) == 1 and bp.get("verified", True)]
        healthy = [bp for bp in bps if bp.get("score", 2) == 2]

        v_pct = (verified / 2.0) * 100
        summary = (
            f"This remediation guide provides detailed, actionable instructions for resolving "
            f"the {len(critical) + len(partial)} issues identified during the Well-Architected Lakehouse "
            f"assessment on {cloud_short}. The verified assessment score is {v_pct:.0f}% based on "
            f"{coverage:.0f}% best practice coverage.\n\n"
            f"Each issue includes an explanation of why it matters, step-by-step remediation instructions, "
            f"and links to the relevant {cloud_short} Databricks documentation."
        )
        doc.add_paragraph(summary)

        doc.add_heading("Issue Summary", level=2)

        table = doc.add_table(rows=4, cols=2)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Count"

        data = [
            ("Critical (Score 0) — Not Implemented", str(len(critical)), _LIGHT_RED_BG),
            ("Needs Improvement (Score 1) — Partial", str(len(partial)), _LIGHT_AMBER_BG),
            ("Healthy (Score 2) — Fully Implemented", str(len(healthy)), _LIGHT_GREEN_BG),
        ]
        for i, (label, count, bg) in enumerate(data):
            row = table.rows[i + 1]
            row.cells[0].text = label
            row.cells[1].text = count
            _shade_cell(row.cells[0], bg)
            _shade_cell(row.cells[1], bg)

        doc.add_paragraph()

        if critical:
            doc.add_heading("Critical Issues (Immediate Action Required)", level=2)
            for bp in critical:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{bp.get('name', 'Unknown')}")
                run.font.bold = True
                run.font.color.rgb = _RED
                p.add_run(f" ({PILLAR_DISPLAY_NAMES.get(bp.get('pillar', ''), bp.get('pillar', ''))})")
                p.add_run(f"\n{bp.get('finding_notes', '')}")

        doc.add_page_break()

    def _add_pillar_section(self, doc: Document, pillar: str, issues: list,
                            cloud: str, cloud_short: str) -> None:
        display = PILLAR_DISPLAY_NAMES.get(pillar, pillar)
        doc.add_heading(f"Pillar: {display}", level=1)

        critical = [bp for bp in issues if bp.get("score", 2) == 0]
        partial = [bp for bp in issues if bp.get("score", 2) == 1 and bp.get("verified", True)]

        if critical:
            p = doc.add_paragraph()
            run = p.add_run(f"{len(critical)} critical issue(s)")
            run.font.bold = True
            run.font.color.rgb = _RED
            p.add_run(f" and {len(partial)} partial implementation(s) found.")
        else:
            p = doc.add_paragraph(f"{len(partial)} partial implementation(s) found. No critical issues.")

        already_covered: set[str] = set()

        for bp in sorted(issues, key=lambda x: x.get("score", 2)):
            score = bp.get("score", 2)
            name = bp.get("name", "Unknown")
            notes = bp.get("finding_notes", "")
            verified = bp.get("verified", True)

            if not verified:
                continue

            score_label = "NOT IMPLEMENTED" if score == 0 else "PARTIAL"
            score_color = _RED if score == 0 else _AMBER

            doc.add_heading(name, level=2)

            p = doc.add_paragraph()
            run = p.add_run(f"Status: {score_label} (Score: {score}/2)")
            run.font.bold = True
            run.font.color.rgb = score_color

            doc.add_heading("Finding", level=3)
            doc.add_paragraph(notes)

            remediation = _find_remediation(name)
            if remediation and remediation["topic"] not in already_covered:
                already_covered.add(remediation["topic"])

                doc.add_heading("Why This Matters", level=3)
                doc.add_paragraph(remediation["why"])

                doc.add_heading("Remediation Steps", level=3)
                for i, step in enumerate(remediation["steps"], 1):
                    p = doc.add_paragraph(style="List Number")
                    p.text = step

                docs_links = _get_docs_for_cloud(remediation, cloud)
                if docs_links:
                    doc.add_heading(f"Documentation ({cloud_short})", level=3)
                    for label, url in docs_links:
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(f"{label}: ")
                        run.font.bold = True
                        link_run = p.add_run(url)
                        link_run.font.color.rgb = _BLUE
                        link_run.font.underline = True
            elif remediation and remediation["topic"] in already_covered:
                p = doc.add_paragraph()
                run = p.add_run(f"See remediation steps under \"{remediation['topic']}\" above.")
                run.font.italic = True
                run.font.color.rgb = _GRAY
            else:
                doc.add_heading("Remediation", level=3)
                doc.add_paragraph(
                    f"Review the Databricks documentation for {cloud_short}-specific guidance "
                    f"on implementing this best practice."
                )
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run("Databricks Best Practices: ")
                run.font.bold = True
                link = p.add_run(_docs_url(cloud, "/getting-started/best-practices"))
                link.font.color.rgb = _BLUE
                link.font.underline = True

    def _add_appendix(self, doc: Document, bps: list, cloud_short: str) -> None:
        doc.add_page_break()
        doc.add_heading("Appendix: All Best Practices", level=1)
        doc.add_paragraph(
            f"Complete list of all {len(bps)} best practices assessed, with scores and verification status."
        )

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Pillar", "Best Practice", "Score", "Verified", "Status"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            _shade_cell(table.rows[0].cells[i], _HEADER_BG)
            for par in table.rows[0].cells[i].paragraphs:
                for run in par.runs:
                    run.font.color.rgb = _WHITE
                    run.font.bold = True

        for bp in bps:
            row = table.add_row()
            score = bp.get("score", 0)
            verified = bp.get("verified", True)
            row.cells[0].text = PILLAR_DISPLAY_NAMES.get(bp.get("pillar", ""), bp.get("pillar", ""))
            row.cells[1].text = bp.get("name", "Unknown")
            row.cells[2].text = str(int(score))
            row.cells[3].text = "Y" if verified else "N"

            if score == 0:
                status = "Not Implemented"
                bg = _LIGHT_RED_BG
            elif score == 1:
                status = "Partial" if verified else "Unverifiable"
                bg = _LIGHT_AMBER_BG if verified else "F5F5F5"
            else:
                status = "Implemented"
                bg = _LIGHT_GREEN_BG
            row.cells[4].text = status
            _shade_cell(row.cells[4], bg)

            for cell in row.cells:
                for par in cell.paragraphs:
                    par.style = doc.styles["Normal"]
                    for run in par.runs:
                        run.font.size = Pt(8)
